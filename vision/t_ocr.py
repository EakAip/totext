#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

# ocr单元测试 / 单线程

# python deepdoc/vision/t_ocr.py --inputs /home/jyd01/wangruihua/ragflow/deepdoc/parser/数据结构C语言版.pdf --output_dir /home/jyd01/wangruihua/ragflow/deepdoc/vision/result

import os
import sys
sys.path.insert(
    0,
    os.path.abspath(
        os.path.join(
            os.path.dirname(
                os.path.abspath(__file__)),
            '../../')))

from seeit import draw_box
# 从当前目录下的__init__.py文件中导入init_in_out函数

from totext.vision import OCR, init_in_out
import argparse
import numpy as np
from tqdm import tqdm
from docx import Document

def save_to_docx(output_dir, page_text):
    docx_file = os.path.join(output_dir, "all_text.docx")
    if not os.path.exists(docx_file):
        doc = Document()
        doc.add_paragraph(page_text)
        doc.save(docx_file)
    else:
        doc = Document(docx_file)
        doc.add_paragraph(page_text)
        doc.save(docx_file)


def main(args):
    ocr = OCR()
    images, outputs = init_in_out(args)
    
    all_text = []

    for i, img in enumerate(tqdm(images, desc="Processing images")):
        bxs = ocr(np.array(img))
        bxs = [(line[0], line[1][0]) for line in bxs]
        bxs = [{
            "text": t,
            "bbox": [b[0][0], b[0][1], b[1][0], b[-1][1]],
            "type": "ocr",
            "score": 1} for b, t in bxs if b[0][0] <= b[1][0] and b[0][1] <= b[-1][1]]
        img = draw_box(images[i], bxs, ["ocr"], 1.)
        img.save(outputs[i], quality=95)
        page_text = "\n".join([o["text"] for o in bxs])
        with open(outputs[i] + ".txt", "w+") as f:
            f.write(page_text)
        
        all_text.append(page_text)

        # 追加到docx文件中
        save_to_docx(args.output_dir, page_text)

    # 将所有页面的文本合并到一个总的txt文件中
    with open(os.path.join(args.output_dir, "all_text.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(all_text))
        


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('--inputs',
                        help="Directory where to store images or PDFs, or a file path to a single image or PDF",
                        required=True)
    parser.add_argument('--output_dir', help="Directory where to store the output images. Default: './ocr_outputs'",
                        default="./ocr_outputs")
    args = parser.parse_args()
    main(args)
