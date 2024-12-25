from PIL import Image
import numpy as np
import os
import sys
import argparse

sys.path.insert(
    0,
    os.path.abspath(
        os.path.join(
            os.path.dirname(
                os.path.abspath(__file__)),
            '../../')))

from vision.seeit import draw_box
from vision import OCR, init_in_out
from tqdm import tqdm
from docx import Document

def save_to_docx(output_dir, page_text):
    docx_file = os.path.join(output_dir, "all_text.docx")
    if not os.path.exists(docx_file):
        doc = Document()
        doc.add_paragraph(page_text)
        doc.save(docx_file)
    else:
        doc = Document(docx_file)
        doc.add_paragraph(page_text)
        doc.save(docx_file)


def main():
    args = argparse.Namespace(inputs='/opt/jyd01/wangruihua/totext/data/测试1.png', output_dir='./ocr_outputs')
    ocr = OCR()
    images, outputs = init_in_out(args)
    
    all_text = []

    for i, img in enumerate(tqdm(images, desc="Processing images")):
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        bxs = ocr(np.array(img))
        bxs = [(line[0], line[1][0]) for line in bxs]
        bxs = [{
            "text": t,
            "bbox": [b[0][0], b[0][1], b[1][0], b[-1][1]],
            "type": "ocr",
            "score": 1} for b, t in bxs if b[0][0] <= b[1][0] and b[0][1] <= b[-1][1]]
        img = draw_box(images[i], bxs, ["ocr"], 1.)
        img.save(outputs[i], quality=95)
        page_text = "\n".join([o["text"] for o in bxs])
        with open(outputs[i] + ".txt", "w+") as f:
            f.write(page_text)
        
        all_text.append(page_text)

        # save_to_docx(args.output_dir, page_text)

    with open(os.path.join(args.output_dir, "all_text.txt"), "w", encoding="utf-8") as f:
        f.write("\n".join(all_text))

    return "\n".join(all_text)


if __name__ == "__main__":
    all_text = main()
    print(all_text)  # 打印出所有识别的文本
